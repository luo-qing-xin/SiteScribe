import io
from datetime import UTC, datetime, timedelta

from docx import Document
from PIL import Image
from pypdf import PdfWriter

from app.auth import hash_password
from app.models import Issue, KnowledgeDocument, ProjectMember, RectificationSubmission, Task, User

SITE_TEXT = "今天3号楼六层钢筋绑扎完成80%，现场12名钢筋工。西侧材料堆放影响通道，通知李班长处理。"


def png_bytes() -> bytes:
    buffer = io.BytesIO()
    Image.new("RGB", (24, 16), "#087354").save(buffer, format="PNG")
    return buffer.getvalue()


def confirmed_issue(auth_client, project, location):
    record = auth_client.post(
        "/api/records",
        json={
            "project_id": project.id,
            "category": "安全巡查",
            "description": SITE_TEXT,
            "occurred_at": datetime.now(UTC).isoformat(),
            "gps_status": "not_requested",
            "location_id": location.id,
        },
    )
    assert record.status_code == 201
    job = auth_client.post(f"/api/records/{record.json()['id']}/event-extractions")
    event_id = auth_client.get(f"/api/event-extraction-jobs/{job.json()['id']}").json()[
        "result_event_id"
    ]
    assert auth_client.post(f"/api/events/{event_id}/confirm").status_code == 200
    issues = auth_client.get(f"/api/projects/{project.id}/issues").json()
    return record.json()["id"], event_id, issues[0]


def test_confirmed_event_syncs_only_safety_issue_idempotently(auth_client, db, project_data):
    project, location, _member = project_data
    record_id, event_id, issue = confirmed_issue(auth_client, project, location)
    assert issue["event_id"] == event_id
    assert issue["record_id"] == record_id
    assert issue["category"] == "文明施工/安全"
    assert issue["description"] == "材料堆放影响通道"
    assert issue["location"] == {"building": "3号楼", "floor": "6层", "zone": "西侧"}
    assert db.query(Issue).filter_by(event_id=event_id).count() == 1
    assert auth_client.post(f"/api/events/{event_id}/confirm").status_code == 200
    assert db.query(Issue).filter_by(event_id=event_id).count() == 1
    assert auth_client.delete(f"/api/records/{record_id}").status_code == 409
    auth_client.post("/api/auth/logout")
    auth_client.post("/api/auth/login", json={"username": "outsider", "password": "demo123"})
    assert auth_client.get(f"/api/issues/{issue['id']}").status_code == 403


def test_knowledge_upload_duplicate_archive_and_project_isolation(
    auth_client, client, db, project_data
):
    project, _location, _member = project_data
    payload = "# 临边防护\n临边防护缺失时应核查并整改。".encode()
    first = auth_client.post(
        f"/api/projects/{project.id}/knowledge-documents",
        files={"file": ("safety.md", payload, "text/markdown")},
    )
    assert first.status_code == 201
    assert first.json()["status"] == "ACTIVE"
    assert first.json()["chunk_count"] == 1
    second = auth_client.post(
        f"/api/projects/{project.id}/knowledge-documents",
        files={"file": ("renamed.md", payload, "text/markdown")},
    )
    assert second.status_code == 201
    assert second.json()["id"] == first.json()["id"]
    assert db.query(KnowledgeDocument).filter_by(project_id=project.id).count() == 1

    other_project = (
        db.query(ProjectMember).filter(ProjectMember.project_id != project.id).first().project_id
    )
    assert auth_client.get(f"/api/projects/{other_project}/knowledge-documents").status_code == 403
    archived = auth_client.post(f"/api/knowledge-documents/{first.json()['id']}/archive")
    assert archived.status_code == 200
    assert archived.json()["status"] == "ARCHIVED"


def test_docx_and_scanned_pdf_parsing(auth_client, project_data):
    project, _location, _member = project_data
    document = Document()
    document.add_heading("高处作业", level=1)
    document.add_paragraph("作业前核查防护设施。")
    docx = io.BytesIO()
    document.save(docx)
    parsed = auth_client.post(
        f"/api/projects/{project.id}/knowledge-documents",
        files={
            "file": (
                "guide.docx",
                docx.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert parsed.status_code == 201
    assert parsed.json()["status"] == "ACTIVE"
    assert parsed.json()["chunk_count"] == 1

    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    pdf = io.BytesIO()
    writer.write(pdf)
    scanned = auth_client.post(
        f"/api/projects/{project.id}/knowledge-documents",
        files={"file": ("scan.pdf", pdf.getvalue(), "application/pdf")},
    )
    assert scanned.status_code == 201
    assert scanned.json()["status"] == "FAILED"
    assert "OCR" in scanned.json()["error_message"]
    assert (
        auth_client.post(f"/api/knowledge-documents/{scanned.json()['id']}/retry").json()["status"]
        == "FAILED"
    )


def test_rag_refuses_empty_library_and_validates_grounded_mock(auth_client, project_data):
    project, location, _member = project_data
    _record_id, _event_id, issue = confirmed_issue(auth_client, project, location)
    empty = auth_client.post(f"/api/issues/{issue['id']}/rag-analyses")
    assert empty.status_code == 201
    assert empty.json()["status"] == "NO_EVIDENCE"
    assert empty.json()["result"] is None

    body = {"seed_demo": (None, "true")}
    seed = auth_client.post(f"/api/projects/{project.id}/knowledge-documents", files=body)
    assert seed.status_code == 201
    assert seed.json()["is_demo"] is True
    result = auth_client.post(f"/api/issues/{issue['id']}/rag-analyses")
    assert result.status_code == 201
    assert result.json()["status"] == "SUCCEEDED"
    citation = result.json()["result"]["citations"][0]
    retrieved = {row["chunk_id"]: row for row in result.json()["retrieved"]}
    assert citation["chunk_id"] in retrieved
    assert citation["excerpt"] in retrieved[citation["chunk_id"]]["content"]
    assert citation["is_demo"] is True
    assert any("非正式演示材料" in warning for warning in result.json()["result"]["warnings"])
    archived = auth_client.post(f"/api/knowledge-documents/{seed.json()['id']}/archive")
    assert archived.json()["status"] == "ARCHIVED"
    historical = auth_client.get(f"/api/rag-analysis-jobs/{result.json()['id']}").json()
    assert historical["result"]["citations"][0] == citation
    assert (
        auth_client.post(f"/api/issues/{issue['id']}/rag-analyses").json()["status"]
        == "NO_EVIDENCE"
    )


def test_rectification_multi_round_review_and_reminders(auth_client, client, db, project_data):
    project, location, member = project_data
    worker = User(
        username="worker",
        password_hash=hash_password("demo123"),
        name="整改责任人",
        role="施工员",
        organization="测试单位",
    )
    db.add(worker)
    db.flush()
    db.add(ProjectMember(project_id=project.id, user_id=worker.id))
    db.commit()
    _record_id, _event_id, issue = confirmed_issue(auth_client, project, location)
    seed_body = {"seed_demo": (None, "true")}
    auth_client.post(f"/api/projects/{project.id}/knowledge-documents", files=seed_body)
    rag = auth_client.post(f"/api/issues/{issue['id']}/rag-analyses").json()
    created = auth_client.post(
        f"/api/issues/{issue['id']}/confirm-and-create-task",
        json={
            "final_description": "材料堆放影响安全通道",
            "rectification_measure": "清理通道并补充整改照片",
            "assignee_id": worker.id,
            "due_at": (datetime.now(UTC) - timedelta(hours=1)).isoformat(),
            "rag_job_id": rag["id"],
        },
    )
    assert created.status_code == 201
    task_id = created.json()["id"]
    duplicate = auth_client.post(
        f"/api/issues/{issue['id']}/confirm-and-create-task",
        json={
            "final_description": "不同描述",
            "rectification_measure": "不同措施",
            "assignee_id": worker.id,
            "due_at": datetime.now(UTC).isoformat(),
            "rag_job_id": rag["id"],
        },
    )
    assert duplicate.json()["id"] == task_id
    assert db.query(Task).filter_by(source_issue_id=issue["id"]).count() == 1

    auth_client.post("/api/auth/logout")
    assert (
        client.post(
            "/api/auth/login", json={"username": "worker", "password": "demo123"}
        ).status_code
        == 200
    )
    assert client.post(f"/api/tasks/{task_id}/start").json()["status"] == "IN_PROGRESS"
    missing = client.post(
        f"/api/tasks/{task_id}/rectification-submissions",
        data={"note": "已清理"},
    )
    assert missing.status_code == 422
    submitted = client.post(
        f"/api/tasks/{task_id}/rectification-submissions",
        data={"note": "第一轮已清理通道"},
        files={"files": ("proof.png", png_bytes(), "image/png")},
    )
    assert submitted.status_code == 201
    assert submitted.json()["status"] == "WAITING_REVIEW"
    proof_id = submitted.json()["submissions"][0]["photos"][0]["id"]
    assert client.delete(f"/api/photos/{proof_id}").status_code == 409
    assert (
        client.post(f"/api/tasks/{task_id}/reviews", json={"decision": "APPROVE"}).status_code
        == 403
    )

    client.post("/api/auth/logout")
    client.post("/api/auth/login", json={"username": "member", "password": "demo123"})
    assert (
        client.post(f"/api/tasks/{task_id}/reviews", json={"decision": "REJECT"}).status_code == 422
    )
    rejected = client.post(
        f"/api/tasks/{task_id}/reviews",
        json={"decision": "REJECT", "reason": "照片未覆盖完整通道"},
    )
    assert rejected.json()["status"] == "IN_PROGRESS"

    client.post("/api/auth/logout")
    client.post("/api/auth/login", json={"username": "worker", "password": "demo123"})
    second = client.post(
        f"/api/tasks/{task_id}/rectification-submissions",
        data={"note": "第二轮补充完整通道照片"},
        files={"files": ("proof2.png", png_bytes(), "image/png")},
    )
    assert second.json()["status"] == "WAITING_REVIEW"
    assert len(second.json()["submissions"]) == 2

    client.post("/api/auth/logout")
    client.post("/api/auth/login", json={"username": "member", "password": "demo123"})
    reminders = client.get(f"/api/projects/{project.id}/task-reminders").json()
    assert task_id in reminders["overdue"]
    assert task_id in reminders["waiting_review"]
    approved = client.post(f"/api/tasks/{task_id}/reviews", json={"decision": "APPROVE"})
    assert approved.json()["status"] == "DONE"
    assert len(approved.json()["submissions"]) == 2
    assert db.query(RectificationSubmission).filter_by(task_id=task_id).count() == 2
    after = client.get(f"/api/projects/{project.id}/task-reminders").json()
    assert task_id not in after["overdue"]
    assert client.delete(f"/api/tasks/{task_id}").status_code == 409

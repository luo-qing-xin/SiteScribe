import hashlib
import io
import re
from dataclasses import dataclass
from pathlib import Path
from uuid import uuid4

from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session

from .config import settings
from .models import KnowledgeChunk, KnowledgeDocument

MAX_KNOWLEDGE_BYTES = 20 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MIME_BY_EXTENSION = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


@dataclass
class Section:
    heading: str | None
    locator: str
    text: str


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text).lower()


def validate_and_read(upload: UploadFile) -> tuple[bytes, str]:
    extension = Path(upload.filename or "").suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=422, detail="仅支持 PDF、DOCX、TXT、MD")
    data = upload.file.read(MAX_KNOWLEDGE_BYTES + 1)
    if not data:
        raise HTTPException(status_code=422, detail="文件不能为空")
    if len(data) > MAX_KNOWLEDGE_BYTES:
        raise HTTPException(status_code=413, detail="单个知识文档不能超过 20MB")
    if extension == ".pdf" and not data.startswith(b"%PDF"):
        raise HTTPException(status_code=422, detail="文件内容不是有效 PDF")
    if extension == ".docx" and not data.startswith(b"PK"):
        raise HTTPException(status_code=422, detail="文件内容不是有效 DOCX")
    return data, extension


def create_document(
    db: Session, project_id: int, user_id: int, upload: UploadFile, title: str | None = None
) -> tuple[KnowledgeDocument, bool]:
    data, extension = validate_and_read(upload)
    digest = hashlib.sha256(data).hexdigest()
    existing = db.query(KnowledgeDocument).filter_by(project_id=project_id, sha256=digest).first()
    if existing:
        return existing, False
    stored_name = f"{uuid4().hex}{extension}"
    settings.knowledge_dir.mkdir(parents=True, exist_ok=True)
    path = settings.knowledge_dir / stored_name
    path.write_bytes(data)
    item = KnowledgeDocument(
        id=str(uuid4()),
        project_id=project_id,
        uploaded_by=user_id,
        title=(title or Path(upload.filename or "文档").stem).strip()[:255],
        original_name=Path(upload.filename or "document").name[:255],
        stored_name=stored_name,
        relative_path=f"knowledge/{stored_name}",
        mime_type=MIME_BY_EXTENSION[extension],
        size_bytes=len(data),
        sha256=digest,
        status="PROCESSING",
        is_demo=False,
    )
    db.add(item)
    db.flush()
    try:
        process_document(db, item, data)
    except Exception as exc:  # parsing failure is persisted and retryable
        item.status = "FAILED"
        item.error_message = str(exc)[:500]
    db.commit()
    return item, True


def process_document(db: Session, document: KnowledgeDocument, data: bytes | None = None) -> None:
    path = settings.knowledge_dir.parent / document.relative_path
    payload = data if data is not None else path.read_bytes()
    extension = Path(document.original_name).suffix.lower()
    sections = extract_sections(payload, extension)
    if not any(section.text.strip() for section in sections):
        raise ValueError("未提取到文字；扫描 PDF 暂不支持 OCR")
    db.query(KnowledgeChunk).filter_by(document_id=document.id).delete()
    chunks = build_chunks(sections)
    for index, (heading, locator, content) in enumerate(chunks):
        db.add(
            KnowledgeChunk(
                document_id=document.id,
                project_id=document.project_id,
                chunk_index=index,
                heading=heading,
                locator=locator,
                content=content,
                normalized_content=normalize_text(content),
            )
        )
    document.status = "ACTIVE"
    document.error_message = None
    db.flush()


def extract_sections(data: bytes, extension: str) -> list[Section]:
    if extension == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return [
            Section(None, f"第 {i} 页", page.extract_text() or "")
            for i, page in enumerate(reader.pages, 1)
        ]
    if extension == ".docx":
        from docx import Document

        document = Document(io.BytesIO(data))
        sections: list[Section] = []
        heading: str | None = None
        for index, paragraph in enumerate(document.paragraphs, 1):
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                heading = text
            else:
                sections.append(Section(heading, f"段落 {index}", text))
        return sections
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError("TXT/MD 必须使用 UTF-8 编码") from exc
    lines = text.splitlines()
    sections = []
    heading = None
    start = 1
    buffer: list[str] = []
    for number, line in enumerate(lines, 1):
        if extension == ".md" and re.match(r"^#{1,6}\s+", line):
            if buffer:
                sections.append(Section(heading, f"第 {start}-{number - 1} 行", "\n".join(buffer)))
            heading, start, buffer = re.sub(r"^#{1,6}\s+", "", line).strip(), number + 1, []
        else:
            buffer.append(line)
    if buffer:
        sections.append(Section(heading, f"第 {start}-{len(lines)} 行", "\n".join(buffer)))
    return sections


def build_chunks(
    sections: list[Section], target: int = 800, overlap: int = 120
) -> list[tuple[str | None, str, str]]:
    output: list[tuple[str | None, str, str]] = []
    for section in sections:
        text = section.text.strip()
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(len(text), start + target)
            if end < len(text):
                boundary = max(text.rfind("\n", start, end), text.rfind("。", start, end))
                if boundary > start + target // 2:
                    end = boundary + 1
            content = text[start:end].strip()
            if content:
                output.append((section.heading, section.locator, content))
            if end >= len(text):
                break
            start = max(start + 1, end - overlap)
    return output
